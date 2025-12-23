import requests
import re 
from pprint import *
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
import os
import shutil

def times14():
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

def clear_folder(folder_path):
    shutil.rmtree(folder_path)
    os.mkdir(folder_path)   






url = input("ВВедите ссылку на статью WIKI:") 
while True:
    if "https://ru.wikipedia.org/wiki/" in url:
        print("Cсылка введeна успешно")
        break
    else:
        print("Неправельная ссылка")

    

headers = {
    "User-Agent":"Mozilla/5.0 (Macintosh; Intel Mac OS X 14_6_0) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36"
}

response = requests.get(url, headers = headers)
response.raise_for_status()
soup = BeautifulSoup(response.text, features="html.parser")

Headline = soup.find("h1")
all_tags = soup.find_all(['p', 'h2',"h3",'h4','h5','img'])


 
doc = Document()
times14()
head = doc.add_heading(Headline.text)
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    

    
for number,delete in enumerate(reversed(all_tags)):
    if "<p" in str(delete):
        print(number)
        break
not_empty_tags = all_tags[0:-number]


os.makedirs("WikiPhoto", exist_ok = True)
clear_folder(folder_path = "WikiPhoto")   
for number,tag in enumerate(not_empty_tags):    
    clear_text = re.sub(r'\[\d*\]', ' ', tag. text)
    if "<p" in str(tag): 
        clear_paragaphs = doc.add_paragraph(clear_text)
        clear_paragaphs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif "<h1" in str(tag):
        head1 = doc.add_heading(clear_text, level = 1)
        head1.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    elif "<h2" in str(tag):
        head2 = doc.add_heading(clear_text, level = 2)
        head2.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    elif '<h3' in str(tag):
        head3 = doc.add_heading(clear_text, level = 3)
        head3.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    elif '<h4' in str(tag):
        head4 = doc.add_heading(clear_text, level = 4)
        head4.alignment = WD_ALIGN_PARAGRAPH.CENTER      
    elif '<h5' in str(tag):
        head5 = doc.add_heading(clear_text, level = 5)
        head5.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    elif "<img" in str(tag):
        img_src = 'https:' + tag["src"]
        response = requests.get(img_src,headers = headers)
        response.raise_for_status()      
        with open(f'WikiPhoto/Photo{number}.png','wb') as file:
           file.write(response.content)
        
        try:
            img2 = doc.add_picture(f'WikiPhoto/Photo{number}.png')
            print(f'WikiPhoto/Photo{number}.png')
            img2 = doc.paragraphs[-1] 
            img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            print(f'WikiPhoto/Photo{number}.png' , "!!!!")
     
doc.save('test.docx')


 